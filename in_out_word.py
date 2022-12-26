import docx
import os


class InOut:

    def __init__(self, orig, transl, name_file):
        self.orig = orig
        self.transl = transl
        self.doc_out = name_file

    def word_write(self,text_out: list)-> None:
        document = docx.Document()
        for var in text_out:
            add_par = document.add_paragraph(var)
            fmt = add_par.paragraph_format
            fmt.space_before = 0
            fmt.space_after = 0
        document.save(self.doc_out)

    @staticmethod
    def list_in(list_len: list, list_str: list, count: int) -> list:
        len_list: list = []
        temp_list: list = []
        summa: int = 0
        for ind, length in enumerate(list_len):
            temp_list.append(ind)
            summa += length
            if summa >= count:
                temp_list.pop()
                len_list.append(temp_list)
                del temp_list
                temp_list = []
                temp_list.append(ind)
                summa = length
        len_list.append(temp_list)
        text_list: list = [".".join(list_str[_].strip() for _ in ind) for ind in len_list]
        return text_list

    def word_read(self) -> tuple:
        base_path = '/home/aleksandr/Downloads/translate'
        # counting 900 characters
        count: int = 900
        list_text: list = []
        len_text: list = []
        os.chdir(base_path)
        file_in = os.listdir(path=".")[0]
        document = docx.Document(file_in)
        pages = document.paragraphs
        lengths_temp: list = [len(paragraph.text) for paragraph in pages]
        for ind, var in enumerate(lengths_temp):
            if var < count:
                list_text.append(pages[ind].text)
                len_text.append(1)
            else:
                string_temp = pages[ind].text.split('.')
                string_len = [len(_) for _ in string_temp]
                out_text: list = self.list_in(string_len, string_temp, count)
                list_text.append(out_text)
                list_count: int = len(out_text)
                len_text.append(list_count)
        text_and_length: tuple = (list_text, len_text,)
        return text_and_length
